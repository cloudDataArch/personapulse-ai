from docx import Document
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.table import WD_TABLE_ALIGNMENT, WD_CELL_VERTICAL_ALIGNMENT
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from docx.shared import Inches, Pt, RGBColor
from pathlib import Path


ROOT = Path(__file__).resolve().parents[1]
OUTPUT = ROOT / "DOCUMENTOS" / "PersonaPulse_AI_Arquitetura_Equipe_Custos_Prototipo.docx"

BLUE = "2E74B5"
DARK_BLUE = "1F4D78"
INK = "0B2545"
LIGHT_BLUE = "E8EEF5"
LIGHT_GRAY = "F2F4F7"
PALE = "F7F9FC"
GREEN = "EAF4EE"

USD_BRL = 5.05


def brl(value):
    return f"R$ {value:,.0f}".replace(",", ".")


def usd_to_brl(value):
    return brl(value * USD_BRL)


def set_cell_shading(cell, fill):
    tc_pr = cell._tc.get_or_add_tcPr()
    shd = tc_pr.find(qn("w:shd"))
    if shd is None:
        shd = OxmlElement("w:shd")
        tc_pr.append(shd)
    shd.set(qn("w:fill"), fill)


def set_cell_margins(cell, top=100, start=120, bottom=100, end=120):
    tc_pr = cell._tc.get_or_add_tcPr()
    tc_mar = tc_pr.first_child_found_in("w:tcMar")
    if tc_mar is None:
        tc_mar = OxmlElement("w:tcMar")
        tc_pr.append(tc_mar)
    for m, v in [("top", top), ("start", start), ("bottom", bottom), ("end", end)]:
        node = tc_mar.find(qn(f"w:{m}"))
        if node is None:
            node = OxmlElement(f"w:{m}")
            tc_mar.append(node)
        node.set(qn("w:w"), str(v))
        node.set(qn("w:type"), "dxa")


def set_table_width(table, widths):
    table.alignment = WD_TABLE_ALIGNMENT.CENTER
    table.autofit = False
    for row in table.rows:
        for idx, width in enumerate(widths):
            cell = row.cells[idx]
            cell.width = Inches(width)
            tc_pr = cell._tc.get_or_add_tcPr()
            tc_w = tc_pr.find(qn("w:tcW"))
            if tc_w is None:
                tc_w = OxmlElement("w:tcW")
                tc_pr.append(tc_w)
            tc_w.set(qn("w:w"), str(int(width * 1440)))
            tc_w.set(qn("w:type"), "dxa")


def style_run(run, bold=False, color=None, size=None):
    run.bold = bold
    if color:
        run.font.color.rgb = RGBColor.from_string(color)
    if size:
        run.font.size = Pt(size)


def add_heading(doc, text, level=1):
    p = doc.add_paragraph(style=f"Heading {level}")
    p.add_run(text)
    return p


def add_body(doc, text):
    p = doc.add_paragraph()
    p.add_run(text)
    return p


def add_bullets(doc, items):
    for item in items:
        p = doc.add_paragraph(style="List Bullet")
        p.add_run(item)


def add_numbered(doc, items):
    for item in items:
        p = doc.add_paragraph(style="List Number")
        p.add_run(item)


def add_callout(doc, title, body, fill=PALE):
    table = doc.add_table(rows=1, cols=1)
    table.style = "Table Grid"
    set_table_width(table, [6.5])
    cell = table.cell(0, 0)
    set_cell_shading(cell, fill)
    set_cell_margins(cell, top=150, bottom=150, start=180, end=180)
    p = cell.paragraphs[0]
    r = p.add_run(title)
    style_run(r, bold=True, color=INK, size=11)
    p.add_run("\n" + body)
    doc.add_paragraph()


def add_table(doc, headers, rows, widths):
    table = doc.add_table(rows=1, cols=len(headers))
    table.style = "Table Grid"
    set_table_width(table, widths)
    for i, h in enumerate(headers):
        cell = table.rows[0].cells[i]
        set_cell_shading(cell, LIGHT_GRAY)
        set_cell_margins(cell)
        cell.vertical_alignment = WD_CELL_VERTICAL_ALIGNMENT.CENTER
        r = cell.paragraphs[0].add_run(h)
        style_run(r, bold=True, color=INK, size=9.5)
    for row in rows:
        cells = table.add_row().cells
        for i, value in enumerate(row):
            set_cell_margins(cells[i])
            cells[i].vertical_alignment = WD_CELL_VERTICAL_ALIGNMENT.CENTER
            r = cells[i].paragraphs[0].add_run(str(value))
            style_run(r, size=9.5)
    doc.add_paragraph()
    return table


def build():
    doc = Document()
    sec = doc.sections[0]
    sec.page_width = Inches(8.5)
    sec.page_height = Inches(11)
    sec.top_margin = Inches(1)
    sec.bottom_margin = Inches(1)
    sec.left_margin = Inches(1)
    sec.right_margin = Inches(1)
    sec.header_distance = Inches(0.492)
    sec.footer_distance = Inches(0.492)

    styles = doc.styles
    normal = styles["Normal"]
    normal.font.name = "Calibri"
    normal.font.size = Pt(11)
    normal.paragraph_format.space_after = Pt(6)
    normal.paragraph_format.line_spacing = 1.10

    for name, size, color, before, after in [
        ("Heading 1", 16, BLUE, 16, 8),
        ("Heading 2", 13, BLUE, 12, 6),
        ("Heading 3", 12, DARK_BLUE, 8, 4),
    ]:
        st = styles[name]
        st.font.name = "Calibri"
        st.font.size = Pt(size)
        st.font.color.rgb = RGBColor.from_string(color)
        st.font.bold = True
        st.paragraph_format.space_before = Pt(before)
        st.paragraph_format.space_after = Pt(after)

    title = doc.add_paragraph()
    title.alignment = WD_ALIGN_PARAGRAPH.CENTER
    r = title.add_run("PersonaPulse AI")
    style_run(r, bold=True, color=INK, size=24)
    subtitle = doc.add_paragraph()
    subtitle.alignment = WD_ALIGN_PARAGRAPH.CENTER
    r = subtitle.add_run("Documento de arquitetura, equipe, ferramentas e custo do protótipo")
    style_run(r, color=DARK_BLUE, size=13)
    meta = doc.add_paragraph()
    meta.alignment = WD_ALIGN_PARAGRAPH.CENTER
    r = meta.add_run("Premissa financeira: 1 USD = R$ 5,05 para estimativas de ferramentas internacionais")
    style_run(r, color="555555", size=10)

    add_callout(
        doc,
        "Resumo executivo",
        "O protótipo recomendado deve ser construído em 10 a 12 semanas, com equipe enxuta e stack cloud gerenciada. O investimento estimado fica entre R$ 260 mil e R$ 330 mil para uma versão apresentável, funcional e testável com dados reais. Uma versão mais enxuta pode cair para R$ 170 mil a R$ 220 mil, reduzindo integrações, QA e profundidade de IA.",
        fill=LIGHT_BLUE,
    )

    add_heading(doc, "1. Arquitetura do Protótipo", 1)
    add_body(doc, "A arquitetura proposta usa serviços gerenciados para acelerar a construção, reduzir necessidade de DevOps pesado e permitir validação comercial antes de investir em infraestrutura enterprise.")

    add_heading(doc, "Fluxo Técnico", 2)
    add_numbered(doc, [
        "Fontes autorizadas: upload CSV, e-commerce ou CRM inicial.",
        "Ingestão: validação de dados, normalização, deduplicação e registro da origem.",
        "Banco operacional: empresas, usuários, clientes, produtos, eventos e campanhas.",
        "Segmentação: regras de negócio para perfis de compra, inatividade, recorrência e intenção.",
        "IA generativa: criação de campanhas, anúncios, e-mails, WhatsApp, SEO e recomendações.",
        "Painel web: dashboard, segmentos, gerador de campanhas, histórico e LGPD.",
        "Observabilidade: logs, erros, uso de IA, eventos do produto e alertas básicos.",
    ])

    architecture_rows = [
        ("Frontend", "Next.js + React + Tailwind + shadcn/ui", "Painel SaaS rápido, moderno e responsivo."),
        ("Backend/API", "FastAPI ou NestJS", "APIs, multiempresa, permissões e regras do produto."),
        ("Banco", "Supabase PostgreSQL + pgvector", "Dados operacionais, autenticação e embeddings iniciais."),
        ("IA", "OpenAI API", "Geração de campanhas, classificação semântica e recomendações."),
        ("Filas/cache", "Upstash Redis", "Processamento assíncrono e cache de tarefas."),
        ("Analytics", "PostHog", "Eventos do produto, funis, uso do painel e comportamento interno."),
        ("Erros/logs", "Sentry", "Rastreamento de erros, performance e alertas de falhas."),
        ("Deploy", "Vercel", "Hospedagem do frontend e deploy contínuo."),
        ("Repositório", "GitHub", "Código, versionamento, issues e CI básico."),
    ]
    add_table(doc, ["Camada", "Ferramenta", "Função no protótipo"], architecture_rows, [1.25, 2.25, 3.00])

    add_heading(doc, "2. Organograma de Construção", 1)
    org_rows = [
        ("Sponsor / Fundador", "1", "Dono da visão, orçamento, validação comercial e decisões finais."),
        ("Product Manager", "0,5 FTE", "Escopo, backlog, roadmap, priorização e aceite das entregas."),
        ("UX/UI Designer", "1", "Protótipo visual, fluxos, telas do painel e design system inicial."),
        ("Tech Lead / Full-stack Senior", "1", "Arquitetura, padrões técnicos, revisão e partes críticas do código."),
        ("Frontend", "1", "Dashboard, telas, estados, gráficos e experiência SaaS."),
        ("Backend", "1", "APIs, autenticação, banco, multiempresa, ingestão e integrações."),
        ("Dados e IA", "0,5 FTE", "Segmentação, prompts, embeddings, scoring e qualidade das respostas."),
        ("QA", "0,5 FTE", "Testes funcionais, regressão, validação de fluxos e checklist de release."),
        ("LGPD / Jurídico", "Consultoria pontual", "Consentimento, privacidade, termos, retenção e exclusão de dados."),
    ]
    add_table(doc, ["Área", "Quantidade", "Responsabilidade"], org_rows, [1.75, 1.20, 3.55])

    add_heading(doc, "3. Ferramentas e Custos Mensais", 1)
    add_body(doc, "Os valores abaixo usam planos adequados para protótipo. O custo real pode variar conforme tráfego, volume de dados, tokens de IA, região, impostos e câmbio.")

    tool_rows = [
        ("Vercel Pro", "3 seats", "US$ 60/mês", usd_to_brl(60), "Deploy e hospedagem do painel."),
        ("Supabase Pro", "1 projeto", "US$ 25/mês", usd_to_brl(25), "Postgres, Auth, Storage e pgvector."),
        ("OpenAI API", "uso controlado", "US$ 150 a US$ 600/mês", f"{usd_to_brl(150)} a {usd_to_brl(600)}", "IA generativa e embeddings."),
        ("Upstash Redis", "Free/PAYG", "US$ 0 a US$ 20/mês", f"{usd_to_brl(0)} a {usd_to_brl(20)}", "Fila leve e cache."),
        ("Sentry Team", "1 time", "US$ 26/mês", usd_to_brl(26), "Erros, performance e alertas."),
        ("PostHog", "free tier inicial", "US$ 0/mês", usd_to_brl(0), "Analytics do produto e funis."),
        ("GitHub", "plano gratuito ou Team", "US$ 0 a US$ 20/mês", f"{usd_to_brl(0)} a {usd_to_brl(20)}", "Repositório e colaboração."),
        ("Domínio/e-mail transacional", "básico", "R$ 50 a R$ 150/mês", "R$ 50 a R$ 150", "Domínio, envio de e-mails e testes."),
    ]
    add_table(doc, ["Ferramenta", "Plano", "USD/mês", "R$/mês", "Uso"], tool_rows, [1.45, 1.15, 1.35, 1.15, 1.40])

    add_callout(
        doc,
        "Estimativa mensal das ferramentas",
        "Faixa enxuta: R$ 1.100 a R$ 2.000/mês. Faixa recomendada, com mais uso de IA: R$ 2.500 a R$ 5.500/mês. Para 3 meses de protótipo, reservar R$ 7 mil a R$ 17 mil para ferramentas é prudente.",
        fill=GREEN,
    )

    add_heading(doc, "4. Profissionais Necessários e Custos", 1)
    add_body(doc, "A tabela considera profissionais no Brasil em regime PJ/freelancer ou contratação por projeto. Os valores são estimativas de mercado para montar um protótipo com qualidade comercial em 10 a 12 semanas.")

    people_rows = [
        ("Product Manager", "0,5", "3 meses", "R$ 18.000/mês", brl(27000), "Escopo, backlog, aceite e rituais."),
        ("UX/UI Designer", "1", "1,5 mês", "R$ 14.000/mês", brl(21000), "Fluxos, protótipo e UI do painel."),
        ("Tech Lead / Full-stack Senior", "1", "3 meses", "R$ 28.000/mês", brl(84000), "Arquitetura e desenvolvimento crítico."),
        ("Frontend React/Next.js", "1", "2,5 meses", "R$ 18.000/mês", brl(45000), "Interface, dashboard e componentes."),
        ("Backend API", "1", "2,5 meses", "R$ 20.000/mês", brl(50000), "APIs, banco, ingestão e integrações."),
        ("Dados e IA", "0,5", "2 meses", "R$ 25.000/mês", brl(25000), "Prompts, scoring e segmentação."),
        ("QA", "0,5", "1,5 mês", "R$ 12.000/mês", brl(9000), "Testes e validação de release."),
        ("LGPD / Jurídico", "consultoria", "30 horas", "R$ 350/hora", brl(10500), "Privacidade, termos e consentimento."),
    ]
    add_table(doc, ["Profissional", "Qtd.", "Duração", "Valor referência", "Subtotal", "Função"], people_rows, [1.45, 0.55, 0.85, 1.10, 0.95, 1.60])

    add_heading(doc, "Resumo de Custo do Protótipo", 2)
    subtotal_people = 271500
    contingency = 40700
    tools_3_months_low = 7000
    tools_3_months_high = 17000
    recommended_low = subtotal_people + contingency + tools_3_months_low
    recommended_high = subtotal_people + contingency + tools_3_months_high
    cost_rows = [
        ("Profissionais", brl(subtotal_people), "Equipe recomendada para 10 a 12 semanas."),
        ("Reserva técnica / gestão de risco", brl(contingency), "Aprox. 15% para retrabalho, ajustes e imprevistos."),
        ("Ferramentas por 3 meses", f"{brl(tools_3_months_low)} a {brl(tools_3_months_high)}", "Cloud, IA, observabilidade e serviços auxiliares."),
        ("Total recomendado", f"{brl(recommended_low)} a {brl(recommended_high)}", "Protótipo funcional e apresentável."),
        ("Total enxuto possível", "R$ 170.000 a R$ 220.000", "Reduzindo equipe, integrações, QA e escopo de IA."),
    ]
    add_table(doc, ["Categoria", "Custo estimado", "Observação"], cost_rows, [2.05, 1.75, 2.70])

    add_heading(doc, "5. O que Entra no Protótipo", 1)
    add_bullets(doc, [
        "Login e estrutura multiempresa básica.",
        "Dashboard com indicadores principais.",
        "Upload de CSV de clientes, produtos e compras.",
        "Segmentação automática com 5 a 7 perfis comportamentais permitidos.",
        "Gerador de campanhas com IA para e-mail, WhatsApp, Meta Ads, Google Ads e SEO.",
        "Recomendações diárias com justificativa.",
        "Histórico de campanhas geradas.",
        "Tela LGPD com origem dos dados, opt-out e exclusão manual.",
        "Monitoramento básico de erros e eventos do produto.",
    ])

    add_heading(doc, "Fora do Protótipo", 2)
    add_bullets(doc, [
        "Integrações complexas com múltiplos CRMs e plataformas ao mesmo tempo.",
        "Data warehouse dedicado em BigQuery ou Snowflake.",
        "Modelos próprios treinados do zero.",
        "Automação completa de mídia paga sem aprovação humana.",
        "Rastreamento externo sem consentimento, câmeras, Wi-Fi público ou fontes sensíveis não autorizadas.",
        "Aplicativo mobile nativo.",
    ])

    add_heading(doc, "6. Cronograma Recomendado", 1)
    timeline_rows = [
        ("Semana 1", "Kickoff, escopo, arquitetura, backlog e definição do nicho inicial."),
        ("Semanas 2-3", "UX/UI, protótipo navegável, modelagem de dados e setup cloud."),
        ("Semanas 4-6", "Base SaaS, autenticação, banco, upload CSV e dashboard inicial."),
        ("Semanas 7-8", "Segmentação, scoring, prompts e gerador de campanhas."),
        ("Semanas 9-10", "Módulo LGPD, histórico, observabilidade e ajustes de experiência."),
        ("Semanas 11-12", "Testes, piloto com dados reais, correções e apresentação comercial."),
    ]
    add_table(doc, ["Período", "Entregas"], timeline_rows, [1.40, 5.10])

    add_heading(doc, "7. Fontes de Preço Consultadas", 1)
    add_bullets(doc, [
        "Vercel Pricing: https://vercel.com/pricing",
        "Supabase Pricing: https://supabase.com/pricing",
        "OpenAI API Pricing: https://openai.com/api/pricing/",
        "Upstash Pricing: https://upstash.com/pricing",
        "Sentry Pricing: https://sentry.io/pricing/",
        "PostHog Pricing: https://posthog.com/pricing",
    ])
    add_body(doc, "Observação: preços de ferramentas internacionais foram consultados em 31/05/2026 e convertidos para reais apenas para estimativa orçamentária. Impostos, IOF, câmbio, uso excedente e contratação anual podem alterar os valores.")

    footer = doc.sections[0].footer.paragraphs[0]
    footer.alignment = WD_ALIGN_PARAGRAPH.CENTER
    r = footer.add_run("PersonaPulse AI - Arquitetura, equipe e custos do protótipo")
    style_run(r, color="777777", size=9)

    doc.save(OUTPUT)


if __name__ == "__main__":
    build()

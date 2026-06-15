from pathlib import Path

from docx import Document
from docx.enum.section import WD_SECTION
from docx.enum.table import WD_TABLE_ALIGNMENT, WD_CELL_VERTICAL_ALIGNMENT
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from docx.shared import Inches, Pt, RGBColor


ROOT = Path(__file__).resolve().parents[1]
OUT_DIR = ROOT / "DOCUMENTOS"
OUT_DIR.mkdir(parents=True, exist_ok=True)
DOCX_PATH = OUT_DIR / "PersonaPulse_AI_Google_Ads_API_Design_Documentation.docx"


BLUE = RGBColor(31, 78, 121)
TEAL = RGBColor(0, 132, 132)
GRAY = RGBColor(89, 89, 89)
LIGHT_BLUE = "EAF3F8"
LIGHT_GRAY = "F2F4F7"


def set_cell_shading(cell, fill):
    tc_pr = cell._tc.get_or_add_tcPr()
    shd = tc_pr.find(qn("w:shd"))
    if shd is None:
        shd = OxmlElement("w:shd")
        tc_pr.append(shd)
    shd.set(qn("w:fill"), fill)


def set_cell_text(cell, text, bold=False, color=None):
    cell.text = ""
    p = cell.paragraphs[0]
    p.paragraph_format.space_after = Pt(0)
    r = p.add_run(text)
    r.bold = bold
    r.font.name = "Calibri"
    r.font.size = Pt(9.5)
    if color:
        r.font.color.rgb = color
    cell.vertical_alignment = WD_CELL_VERTICAL_ALIGNMENT.CENTER


def add_kv_table(doc, rows):
    table = doc.add_table(rows=0, cols=2)
    table.alignment = WD_TABLE_ALIGNMENT.CENTER
    table.style = "Table Grid"
    widths = [Inches(1.8), Inches(4.7)]
    for label, value in rows:
        cells = table.add_row().cells
        for i, width in enumerate(widths):
            cells[i].width = width
        set_cell_shading(cells[0], LIGHT_GRAY)
        set_cell_text(cells[0], label, bold=True, color=BLUE)
        set_cell_text(cells[1], value)
    doc.add_paragraph()
    return table


def add_matrix(doc, headers, rows):
    table = doc.add_table(rows=1, cols=len(headers))
    table.alignment = WD_TABLE_ALIGNMENT.CENTER
    table.style = "Table Grid"
    for cell, header in zip(table.rows[0].cells, headers):
        set_cell_shading(cell, LIGHT_BLUE)
        set_cell_text(cell, header, bold=True, color=BLUE)
    for row in rows:
        cells = table.add_row().cells
        for cell, value in zip(cells, row):
            set_cell_text(cell, value)
    doc.add_paragraph()
    return table


def add_heading(doc, text, level=1):
    p = doc.add_heading(text, level=level)
    for run in p.runs:
        run.font.name = "Calibri"
        run.font.color.rgb = BLUE if level < 3 else TEAL
    return p


def add_para(doc, text, bold_prefix=None):
    p = doc.add_paragraph()
    p.paragraph_format.space_after = Pt(6)
    if bold_prefix and text.startswith(bold_prefix):
        r1 = p.add_run(bold_prefix)
        r1.bold = True
        r1.font.name = "Calibri"
        r1.font.size = Pt(10.5)
        r2 = p.add_run(text[len(bold_prefix):])
        r2.font.name = "Calibri"
        r2.font.size = Pt(10.5)
    else:
        r = p.add_run(text)
        r.font.name = "Calibri"
        r.font.size = Pt(10.5)
    return p


def add_bullets(doc, items):
    for item in items:
        p = doc.add_paragraph(style="List Bullet")
        p.paragraph_format.space_after = Pt(3)
        r = p.add_run(item)
        r.font.name = "Calibri"
        r.font.size = Pt(10)


def add_numbered(doc, items):
    for item in items:
        p = doc.add_paragraph(style="List Number")
        p.paragraph_format.space_after = Pt(3)
        r = p.add_run(item)
        r.font.name = "Calibri"
        r.font.size = Pt(10)


def build_doc():
    doc = Document()
    section = doc.sections[0]
    section.top_margin = Inches(0.8)
    section.bottom_margin = Inches(0.8)
    section.left_margin = Inches(0.8)
    section.right_margin = Inches(0.8)

    styles = doc.styles
    styles["Normal"].font.name = "Calibri"
    styles["Normal"].font.size = Pt(10.5)

    title = doc.add_paragraph()
    title.alignment = WD_ALIGN_PARAGRAPH.LEFT
    run = title.add_run("PersonaPulse AI\nGoogle Ads API Design Documentation")
    run.font.name = "Calibri"
    run.font.size = Pt(24)
    run.bold = True
    run.font.color.rgb = BLUE

    subtitle = doc.add_paragraph()
    subtitle.paragraph_format.space_after = Pt(12)
    r = subtitle.add_run(
        "Prepared for Google Ads API Developer Token Basic/Standard Access Review"
    )
    r.font.name = "Calibri"
    r.font.size = Pt(12)
    r.font.color.rgb = GRAY

    add_kv_table(
        doc,
        [
            ("Company", "COBWEB ARQUITETURA E CONSULTORIA DADOS"),
            ("Product", "PersonaPulse AI"),
            ("Primary URL", "https://personapulse-ai.onrender.com/app"),
            ("API Base URL", "https://personapulse-ai.onrender.com"),
            ("Contact email", "cloud.datascience.arch@gmail.com"),
            ("Document purpose", "Describe the design, data flow, permissions, security controls and Google Ads API usage of the PersonaPulse AI tool."),
        ],
    )

    add_heading(doc, "1. Product Overview")
    add_para(
        doc,
        "PersonaPulse AI is a marketing intelligence and campaign performance platform. It helps authorized business users import consented customer data from CRM or CSV, connect advertising platforms, generate campaign drafts, monitor campaign performance, and analyze executive indicators such as conversions, attributed revenue, spend, ROAS and ROI.",
    )
    add_para(
        doc,
        "The Google Ads API is used to retrieve campaign and performance data from Google Ads accounts that the user is authorized to access through OAuth. The retrieved information is displayed in dashboards and may be exported to executive reporting tools such as Power BI.",
    )

    add_heading(doc, "2. Intended Google Ads API Use")
    add_bullets(
        doc,
        [
            "Read campaign metadata such as campaign ID, campaign name, status and advertising channel type.",
            "Read aggregated campaign metrics such as impressions, clicks, conversions, cost and conversion value.",
            "Calculate business indicators inside PersonaPulse AI, including ROI, ROAS, cost over revenue and campaign performance comparisons.",
            "Display campaign-level performance to authorized users in the PersonaPulse AI dashboard.",
            "Support Power BI reporting through dedicated API endpoints that expose summarized and tabular data.",
        ],
    )
    add_para(
        doc,
        "The current MVP focuses on reading and analyzing authorized campaign data. It does not scrape Google Ads interfaces, does not bypass Google Ads permissions, and does not access Google user-level personal data.",
    )

    add_heading(doc, "3. System Architecture")
    add_matrix(
        doc,
        ["Layer", "Component", "Purpose"],
        [
            ("Frontend", "PersonaPulse AI Web App", "Browser-based interface served from the backend at /app. Users import data, configure integrations, generate campaigns and view dashboards."),
            ("Backend API", "Python HTTP service", "Receives CRM data, handles OAuth callbacks, synchronizes Ads data, calculates indicators and exposes Power BI endpoints."),
            ("Persistence", "PostgreSQL", "Stores customers, orders, events, campaigns, campaign metrics, recommendations, connector configurations, audit records and price research results."),
            ("External API", "Google Ads API", "Retrieves authorized campaign and metric data using OAuth and developer token access."),
            ("External API", "Google Custom Search API", "Retrieves market price references for product research where configured."),
            ("Reporting", "Power BI endpoints", "Provides JSON endpoints for executive summary, customers, campaigns and data sources."),
        ],
    )

    add_heading(doc, "4. High-Level Data Flow")
    add_numbered(
        doc,
        [
            "The user signs in to the relevant Google account and authorizes PersonaPulse AI through the Google OAuth consent screen.",
            "Google redirects the user to PersonaPulse AI's OAuth callback endpoint with a temporary authorization code and state value.",
            "The backend validates the state value, exchanges the code for access/refresh tokens, and stores token data server-side.",
            "When the user clicks Synchronize, the backend calls the Google Ads API with the configured developer token, customer ID and optional MCC login customer ID.",
            "The backend normalizes the Google Ads response into internal campaign and campaign metric records.",
            "PersonaPulse AI updates dashboards, campaign cards, ROI/ROAS metrics and Power BI output endpoints.",
        ],
    )

    add_heading(doc, "5. OAuth and Authorization Design")
    add_kv_table(
        doc,
        [
            ("OAuth provider", "Google OAuth 2.0"),
            ("Requested scope", "https://www.googleapis.com/auth/adwords"),
            ("Redirect URI", "https://personapulse-ai.onrender.com/api/oauth/callback"),
            ("Authorization protection", "State parameter validation is used to reduce CSRF and callback misuse risk."),
            ("Token handling", "OAuth tokens are handled by the backend service and are not exposed in the browser UI."),
            ("Account access", "The system can only synchronize customer accounts that the authorized Google user is allowed to access."),
        ],
    )
    add_para(
        doc,
        "If a manager account (MCC) is used, PersonaPulse AI uses the MCC login customer ID to discover or access client accounts and retrieves metrics from the client account level, not from the manager account directly.",
    )

    add_heading(doc, "6. Data Collected from Google Ads API")
    add_matrix(
        doc,
        ["Data group", "Fields", "Use inside PersonaPulse AI"],
        [
            ("Campaign metadata", "Campaign ID, name, status, channel type", "Identify campaigns, display campaign cards and allow users to select campaign-level dashboards."),
            ("Campaign metrics", "Impressions, clicks, conversions, cost, conversion value", "Calculate performance, ROI, ROAS, conversion rate and spend analysis."),
            ("Account identifiers", "Customer ID and optional MCC login customer ID", "Route Google Ads API requests to the authorized account."),
        ],
    )
    add_para(
        doc,
        "PersonaPulse AI does not request or store Google Ads end-user personal data. The product uses aggregated account and campaign-level information for business analytics.",
    )

    add_heading(doc, "7. Data Stored by PersonaPulse AI")
    add_matrix(
        doc,
        ["Table / domain", "Examples of stored information"],
        [
            ("customers", "Customer records imported from authorized CRM or CSV sources, including source labels and consent flags."),
            ("orders", "Purchase records and attributed revenue from CRM or imported sources."),
            ("events", "Consent-based behavioral events received from connected business systems."),
            ("campaigns", "Campaign records generated in PersonaPulse or imported from Ads platforms."),
            ("campaign_metrics", "Campaign-level performance metrics from Ads platforms."),
            ("connector_configs", "Connector configuration records needed for authorized integrations."),
            ("audit", "Operational audit events such as sync, import, OAuth and deletion events."),
            ("price_researches", "Structured market price research results from configured pricing APIs."),
        ],
    )

    add_heading(doc, "8. Security and Privacy Controls")
    add_bullets(
        doc,
        [
            "Production traffic is served over HTTPS through Render.",
            "OAuth uses a redirect URI registered in Google Cloud and a state parameter for callback validation.",
            "Secrets and production credentials are intended to be stored as backend environment variables or server-side configuration, not committed to source code.",
            "The application separates data sources such as CSV, CRM, Meta Ads, Google Ads and other Ads connectors to avoid accidental data loss or mixing.",
            "The interface includes deletion controls by data source and audit records for important operational actions.",
            "Customer data used for marketing analysis is expected to be consented and governed by LGPD/privacy policies.",
            "Google Ads data is used only for authorized campaign analysis and reporting, not for unauthorized profiling or scraping.",
        ],
    )

    add_heading(doc, "9. User Access and Roles")
    add_para(
        doc,
        "The MVP is designed for authorized business users who manage marketing, CRM, campaign performance and executive reporting. In production, the intended access model is role-based: administrators configure integrations, analysts review dashboards and campaigns, and executive users consume summarized reports.",
    )
    add_para(
        doc,
        "External customer access may be offered only to clients who have authorized their own data sources and advertising accounts. The tool does not grant access to accounts or data unless the user has completed the required authorization flow.",
    )

    add_heading(doc, "10. Google Ads API Compliance Notes")
    add_bullets(
        doc,
        [
            "PersonaPulse AI will not share, sell or expose Google Ads API data outside the authorized user's workspace and reporting outputs.",
            "The application will not scrape Google Ads web interfaces or use automated browser extraction to obtain Ads data.",
            "The application will respect Google Ads API permissions, rate limits and token approval level restrictions.",
            "Users must authorize access through OAuth and may revoke access from their Google account settings.",
            "If write operations are added in the future, they will be clearly separated from read-only reporting and will require explicit user confirmation.",
        ],
    )

    add_heading(doc, "11. Current MVP Status")
    add_matrix(
        doc,
        ["Capability", "Status"],
        [
            ("Web dashboard", "Implemented"),
            ("CRM synchronization API", "Implemented"),
            ("PostgreSQL persistence", "Implemented"),
            ("Power BI endpoints", "Implemented"),
            ("Google Ads OAuth configuration", "Implemented"),
            ("Google Ads real sync", "Implemented, pending Google Developer Token approval for non-test accounts"),
            ("Google CSE product price research", "Implemented when credentials are configured"),
        ],
    )

    add_heading(doc, "12. Review Request")
    add_para(
        doc,
        "COBWEB ARQUITETURA E CONSULTORIA DADOS requests Google Ads API access for PersonaPulse AI so the platform can read authorized Google Ads campaign data from real advertiser accounts and display performance analytics inside the product. The requested access is necessary because test-only developer token access prevents the product from validating real campaign dashboards, ROI analysis and executive reporting with authorized business accounts.",
    )

    doc.add_section(WD_SECTION.NEW_PAGE)
    add_heading(doc, "Appendix A - Endpoint Summary")
    add_matrix(
        doc,
        ["Endpoint", "Purpose"],
        [
            ("/health", "Service and persistence health check."),
            ("/app", "PersonaPulse AI web application."),
            ("/api/oauth/start?source=google", "Starts Google OAuth authorization."),
            ("/api/oauth/callback", "Receives OAuth callback and exchanges authorization code."),
            ("/api/connectors/google/sync", "Synchronizes authorized Google Ads campaign data."),
            ("/api/powerbi/executive-summary", "Returns executive summary for Power BI."),
            ("/api/powerbi/customers", "Returns customer records for Power BI."),
            ("/api/powerbi/campaigns", "Returns campaign records for Power BI."),
            ("/api/powerbi/sources", "Returns data source summary for Power BI."),
        ],
    )

    footer = doc.sections[0].footer.paragraphs[0]
    footer.alignment = WD_ALIGN_PARAGRAPH.CENTER
    fr = footer.add_run("PersonaPulse AI - Google Ads API Design Documentation")
    fr.font.size = Pt(8)
    fr.font.color.rgb = GRAY

    doc.save(DOCX_PATH)
    return DOCX_PATH


if __name__ == "__main__":
    print(build_doc())

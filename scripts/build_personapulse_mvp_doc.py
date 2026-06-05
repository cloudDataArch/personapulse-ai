from docx import Document
from docx.enum.section import WD_SECTION
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.table import WD_TABLE_ALIGNMENT, WD_CELL_VERTICAL_ALIGNMENT
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from docx.shared import Inches, Pt, RGBColor
from pathlib import Path


ROOT = Path(__file__).resolve().parents[1]
OUTPUT = ROOT / "DOCUMENTOS" / "PersonaPulse_AI_Organograma_Etapas_MVP.docx"


BLUE = "2E74B5"
DARK_BLUE = "1F4D78"
INK = "0B2545"
LIGHT_BLUE = "E8EEF5"
LIGHT_GRAY = "F2F4F7"
PALE = "F7F9FC"
WHITE = "FFFFFF"


def set_cell_shading(cell, fill):
    tc_pr = cell._tc.get_or_add_tcPr()
    shd = tc_pr.find(qn("w:shd"))
    if shd is None:
        shd = OxmlElement("w:shd")
        tc_pr.append(shd)
    shd.set(qn("w:fill"), fill)


def set_cell_margins(cell, top=100, start=120, bottom=100, end=120):
    tc = cell._tc
    tc_pr = tc.get_or_add_tcPr()
    tc_mar = tc_pr.first_child_found_in("w:tcMar")
    if tc_mar is None:
        tc_mar = OxmlElement("w:tcMar")
        tc_pr.append(tc_mar)
    for m, v in [("top", top), ("start", start), ("bottom", bottom), ("end", end)]:
        node = tc_mar.find(qn(f"w:{m}"))
        if node is None:
            node = OxmlElement(f"w:{m}")
            tc_mar.append(node)
        node.set(qn("w:w"), str(v))
        node.set(qn("w:type"), "dxa")


def set_table_width(table, widths):
    table.alignment = WD_TABLE_ALIGNMENT.CENTER
    table.autofit = False
    for row in table.rows:
        for idx, width in enumerate(widths):
            cell = row.cells[idx]
            cell.width = Inches(width)
            tc_pr = cell._tc.get_or_add_tcPr()
            tc_w = tc_pr.find(qn("w:tcW"))
            if tc_w is None:
                tc_w = OxmlElement("w:tcW")
                tc_pr.append(tc_w)
            tc_w.set(qn("w:w"), str(int(width * 1440)))
            tc_w.set(qn("w:type"), "dxa")


def style_run(run, bold=False, color=None, size=None):
    run.bold = bold
    if color:
        run.font.color.rgb = RGBColor.from_string(color)
    if size:
        run.font.size = Pt(size)


def add_heading(doc, text, level=1):
    p = doc.add_paragraph()
    p.style = f"Heading {level}"
    p.add_run(text)
    return p


def add_body(doc, text, bold_prefix=None):
    p = doc.add_paragraph()
    if bold_prefix and text.startswith(bold_prefix):
        r = p.add_run(bold_prefix)
        r.bold = True
        p.add_run(text[len(bold_prefix):])
    else:
        p.add_run(text)
    return p


def add_bullets(doc, items):
    for item in items:
        p = doc.add_paragraph(style="List Bullet")
        p.add_run(item)


def add_numbered(doc, items):
    for item in items:
        p = doc.add_paragraph(style="List Number")
        p.add_run(item)


def add_callout(doc, title, body, fill=PALE):
    table = doc.add_table(rows=1, cols=1)
    table.style = "Table Grid"
    set_table_width(table, [6.5])
    cell = table.cell(0, 0)
    set_cell_shading(cell, fill)
    set_cell_margins(cell, top=140, bottom=140, start=180, end=180)
    p = cell.paragraphs[0]
    r = p.add_run(title)
    style_run(r, bold=True, color=INK, size=11)
    p.add_run("\n" + body)
    doc.add_paragraph()


def add_table(doc, headers, rows, widths):
    table = doc.add_table(rows=1, cols=len(headers))
    table.style = "Table Grid"
    set_table_width(table, widths)
    hdr = table.rows[0].cells
    for i, h in enumerate(headers):
        set_cell_shading(hdr[i], LIGHT_GRAY)
        set_cell_margins(hdr[i])
        hdr[i].vertical_alignment = WD_CELL_VERTICAL_ALIGNMENT.CENTER
        p = hdr[i].paragraphs[0]
        r = p.add_run(h)
        style_run(r, bold=True, color=INK, size=10)
    for row in rows:
        cells = table.add_row().cells
        for i, value in enumerate(row):
            set_cell_margins(cells[i])
            cells[i].vertical_alignment = WD_CELL_VERTICAL_ALIGNMENT.CENTER
            p = cells[i].paragraphs[0]
            r = p.add_run(value)
            style_run(r, size=10)
    doc.add_paragraph()
    return table


def build():
    doc = Document()
    sec = doc.sections[0]
    sec.page_width = Inches(8.5)
    sec.page_height = Inches(11)
    sec.top_margin = Inches(1)
    sec.bottom_margin = Inches(1)
    sec.left_margin = Inches(1)
    sec.right_margin = Inches(1)
    sec.header_distance = Inches(0.492)
    sec.footer_distance = Inches(0.492)

    styles = doc.styles
    normal = styles["Normal"]
    normal.font.name = "Calibri"
    normal.font.size = Pt(11)
    normal.paragraph_format.space_after = Pt(6)
    normal.paragraph_format.line_spacing = 1.10

    for name, size, color, before, after in [
        ("Heading 1", 16, BLUE, 16, 8),
        ("Heading 2", 13, BLUE, 12, 6),
        ("Heading 3", 12, DARK_BLUE, 8, 4),
    ]:
        st = styles[name]
        st.font.name = "Calibri"
        st.font.size = Pt(size)
        st.font.color.rgb = RGBColor.from_string(color)
        st.font.bold = True
        st.paragraph_format.space_before = Pt(before)
        st.paragraph_format.space_after = Pt(after)

    title = doc.add_paragraph()
    title.alignment = WD_ALIGN_PARAGRAPH.CENTER
    r = title.add_run("PersonaPulse AI")
    style_run(r, bold=True, color=INK, size=24)
    subtitle = doc.add_paragraph()
    subtitle.alignment = WD_ALIGN_PARAGRAPH.CENTER
    r = subtitle.add_run("Organograma de construção, etapas de execução e escopo do MVP")
    style_run(r, color=DARK_BLUE, size=13)
    meta = doc.add_paragraph()
    meta.alignment = WD_ALIGN_PARAGRAPH.CENTER
    r = meta.add_run("Documento executivo para planejamento técnico e construção inicial")
    style_run(r, color="555555", size=10)

    add_callout(
        doc,
        "Objetivo do MVP",
        "Construir uma plataforma SaaS de inteligência comportamental consentida que permita importar dados próprios, segmentar clientes, gerar campanhas com IA generativa e acompanhar recomendações diárias com governança LGPD.",
        fill=LIGHT_BLUE,
    )

    add_heading(doc, "1. Organograma de Construção da Ferramenta", 1)
    add_body(
        doc,
        "O organograma abaixo organiza as frentes necessárias para construir o PersonaPulse AI. Ele mostra quem decide, quem constrói e quais blocos precisam caminhar juntos para que o MVP chegue ao mercado com qualidade.",
    )

    org_rows = [
        ("Fundador / Sponsor", "Define visão, prioridades, orçamento, metas comerciais e validação com clientes.", "Decisões de produto, investimento e go-to-market."),
        ("Product Manager", "Transforma a visão em backlog, escopo, roadmap, critérios de aceite e priorização.", "Documento de requisitos, backlog e plano de releases."),
        ("UX/UI Designer", "Desenha fluxo do painel, telas, jornadas e experiência de uso.", "Protótipo navegável e design system inicial."),
        ("Tech Lead", "Define arquitetura, padrões técnicos, segurança e integração entre times.", "Arquitetura, decisões técnicas e revisão de implementação."),
        ("Frontend", "Constrói o painel web, dashboard, telas de segmentos, campanhas e privacidade.", "Aplicação Next.js/React funcional."),
        ("Backend", "Constrói APIs, autenticação, multiempresa, regras de negócio e integrações.", "API SaaS, serviços internos e camada de permissões."),
        ("Dados e IA", "Cria segmentação, scoring, embeddings, prompts e geração de campanhas.", "Motor de recomendação e geração com IA."),
        ("DevOps / Cloud", "Configura cloud, deploy, monitoramento, banco, filas e ambientes.", "Ambiente de produção, staging, logs e backups."),
        ("LGPD / Jurídico", "Define consentimento, política de privacidade, retenção e direitos do titular.", "Modelo de governança, textos legais e requisitos de compliance."),
        ("QA / Testes", "Valida fluxos, segurança básica, regressões, performance e qualidade de dados.", "Plano de testes e aprovação de release."),
    ]
    add_table(doc, ["Frente", "Responsabilidade", "Entregável"], org_rows, [1.65, 2.95, 1.90])

    add_heading(doc, "Organograma Operacional Simplificado", 2)
    add_body(doc, "Fundador / Sponsor")
    add_bullets(doc, [
        "Product Manager: organiza escopo, roadmap e validação com clientes.",
        "Tech Lead: lidera arquitetura e decisões técnicas.",
        "UX/UI Designer: define experiência, telas e protótipo.",
        "LGPD / Jurídico: garante consentimento, privacidade e governança.",
    ])
    add_body(doc, "Tech Lead")
    add_bullets(doc, [
        "Frontend: painel web e experiência do usuário.",
        "Backend: APIs, autenticação, multiempresa e integrações.",
        "Dados e IA: segmentação, scoring, prompts e geração de campanhas.",
        "DevOps / Cloud: deploy, infraestrutura, observabilidade e backups.",
        "QA / Testes: validação funcional, técnica e de segurança.",
    ])

    add_heading(doc, "2. Etapas que Deverão ser Seguidas", 1)
    add_body(
        doc,
        "A construção deve seguir uma ordem prática: validar o problema, desenhar o produto, construir a base técnica, adicionar inteligência, testar com dados reais e preparar a primeira venda.",
    )
    steps = [
        ("1", "Descoberta e validação", "Entrevistar 5 a 10 empresas-alvo, escolher nicho inicial e confirmar dores de marketing, dados e campanhas.", "Problema validado, persona compradora e proposta de valor."),
        ("2", "Definição do MVP", "Priorizar apenas as funções essenciais: importar dados, segmentar, gerar campanhas, recomendar ações e controlar consentimento.", "Escopo fechado do MVP e backlog inicial."),
        ("3", "Protótipo UX/UI", "Criar telas do dashboard, segmentos, gerador de campanhas, recomendações e privacidade.", "Protótipo navegável aprovado."),
        ("4", "Arquitetura técnica", "Definir cloud, banco, APIs, filas, modelo de dados, autenticação e camadas de IA.", "Documento técnico e plano de infraestrutura."),
        ("5", "Base SaaS", "Construir login, organizações, usuários, permissões, projetos e estrutura multiempresa.", "Aplicação base funcionando."),
        ("6", "Importação de dados", "Permitir upload de CSV e integração inicial com e-commerce ou CRM.", "Dados entrando no sistema com validação."),
        ("7", "Segmentação e scoring", "Criar regras para clientes recorrentes, inativos, alto valor, desconto, abandono e recompra.", "Segmentos automáticos e pontuação de intenção."),
        ("8", "IA generativa", "Gerar campanhas, anúncios, e-mails, WhatsApp, SEO e explicações das recomendações.", "Motor de geração e biblioteca de prompts."),
        ("9", "Painel e recomendações", "Exibir oportunidades diárias, indicadores, campanhas sugeridas e resultados.", "Dashboard operacional do MVP."),
        ("10", "LGPD e auditoria", "Implementar consentimento, opt-out, exclusão, origem dos dados e logs básicos.", "Módulo de privacidade e trilha de auditoria."),
        ("11", "Testes e piloto", "Rodar com 1 a 3 clientes reais, medir qualidade dos dados, campanhas e usabilidade.", "Relatório de piloto e ajustes finais."),
        ("12", "Lançamento comercial", "Criar pricing, onboarding, materiais de venda e suporte inicial.", "MVP pronto para venda assistida."),
    ]
    add_table(doc, ["Etapa", "Nome", "Ação principal", "Resultado esperado"], steps, [0.55, 1.55, 2.95, 1.45])

    add_heading(doc, "3. MVP Consolidado", 1)
    add_body(
        doc,
        "O MVP deve provar uma tese simples: com dados próprios e consentidos, a empresa consegue conhecer melhor seus clientes, gerar campanhas melhores e tomar decisões diárias com mais precisão.",
    )

    add_heading(doc, "Escopo Funcional do MVP", 2)
    features = [
        ("Dashboard", "Visão diária de clientes, segmentos, campanhas geradas, recomendações e conversões.", "Obrigatório"),
        ("Upload de CSV", "Importação de clientes, pedidos, produtos e eventos básicos.", "Obrigatório"),
        ("Integração inicial", "Escolher uma: Shopify, WooCommerce, HubSpot, RD Station ou Google Analytics.", "Obrigatório"),
        ("Segmentação automática", "Classificar clientes por recorrência, valor, inatividade, desconto, abandono e intenção.", "Obrigatório"),
        ("Gerador de campanhas", "Criar textos para e-mail, WhatsApp, Meta Ads, Google Ads e SEO.", "Obrigatório"),
        ("Recomendações diárias", "Sugerir público, produto, mensagem, canal e justificativa.", "Obrigatório"),
        ("Histórico de campanhas", "Guardar campanhas geradas, status e resultado informado.", "Obrigatório"),
        ("Módulo LGPD", "Consentimento, origem dos dados, opt-out, exclusão e logs.", "Obrigatório"),
        ("Admin SaaS", "Empresas, usuários, permissões e cobrança futura.", "Desejável no MVP"),
        ("Relatório executivo", "Resumo semanal com ganhos, riscos e próximos testes.", "Desejável no MVP"),
    ]
    add_table(doc, ["Módulo", "Descrição", "Prioridade"], features, [1.55, 3.85, 1.10])

    add_heading(doc, "Stack Recomendada para o MVP", 2)
    stack = [
        ("Frontend", "Next.js, React, Tailwind CSS, shadcn/ui, Recharts", "Painel web rápido, moderno e fácil de evoluir."),
        ("Backend", "FastAPI ou NestJS", "APIs, regras de negócio, autenticação e integrações."),
        ("Banco", "Supabase PostgreSQL com pgvector", "Dados operacionais, embeddings e busca semântica inicial."),
        ("IA", "OpenAI API", "Geração de campanhas, análise de perfis, SEO e recomendações."),
        ("Filas/cache", "Upstash Redis ou Redis gerenciado", "Processamento assíncrono e cache."),
        ("Cloud inicial", "Vercel + Supabase + OpenAI", "Menor tempo de construção e menor custo inicial."),
        ("Monitoramento", "Sentry + PostHog", "Erros, eventos, funil e comportamento dentro do produto."),
        ("Pagamentos", "Stripe ou Mercado Pago", "Assinatura SaaS e cobrança recorrente."),
    ]
    add_table(doc, ["Camada", "Ferramentas", "Motivo"], stack, [1.35, 2.55, 2.60])

    add_heading(doc, "Arquitetura do MVP", 2)
    architecture = [
        "Fontes de dados: CSV, e-commerce, CRM ou analytics autorizado.",
        "Camada de ingestão: valida dados, remove duplicidades e registra origem.",
        "Banco operacional: guarda empresas, usuários, clientes, eventos, produtos e campanhas.",
        "Motor de segmentação: cria grupos e scores com regras claras.",
        "Motor de IA: transforma segmentos em campanhas, SEO, anúncios e recomendações.",
        "Painel web: exibe dashboard, segmentos, campanhas, recomendações e privacidade.",
        "Módulo LGPD: controla consentimento, opt-out, exclusão e auditoria.",
    ]
    add_numbered(doc, architecture)

    add_heading(doc, "Critérios de Pronto para o MVP", 2)
    readiness = [
        "Uma empresa consegue criar conta, importar dados e visualizar clientes.",
        "O sistema cria pelo menos 6 segmentos automáticos úteis.",
        "A IA gera campanhas em pelo menos 5 formatos: e-mail, WhatsApp, Meta Ads, Google Ads e SEO.",
        "O dashboard mostra recomendações diárias com justificativa.",
        "O módulo LGPD registra origem dos dados e permite exclusão ou opt-out.",
        "O produto funciona para pelo menos 1 nicho inicial, preferencialmente e-commerce.",
        "O piloto com clientes reais gera métricas de uso e feedback comercial.",
    ]
    add_bullets(doc, readiness)

    add_heading(doc, "Cronograma Sugerido", 2)
    timeline = [
        ("Semanas 1-2", "Descoberta, nicho inicial, requisitos e escopo do MVP."),
        ("Semanas 3-4", "Protótipo UX/UI, arquitetura técnica e setup de cloud."),
        ("Semanas 5-8", "Base SaaS, autenticação, banco, upload de CSV e primeiras telas."),
        ("Semanas 9-12", "Segmentação, IA generativa, recomendações e histórico de campanhas."),
        ("Semanas 13-16", "Módulo LGPD, testes, piloto com clientes e ajustes finais."),
        ("Semanas 17-20", "Lançamento comercial assistido, onboarding e melhoria por feedback."),
    ]
    add_table(doc, ["Período", "Entregas principais"], timeline, [1.35, 5.15])

    add_heading(doc, "Estimativa de Construção", 2)
    add_callout(
        doc,
        "Faixa recomendada para iniciar",
        "Para construir um MVP comercial enxuto, a estimativa realista fica entre R$ 120 mil e R$ 300 mil, com prazo médio de 3 a 5 meses. A cloud inicial tende a ficar entre R$ 1.500 e R$ 8.000 por mês, variando conforme volume de dados e uso de IA.",
        fill=LIGHT_GRAY,
    )

    add_heading(doc, "Próxima Decisão", 1)
    add_body(
        doc,
        "A decisão mais importante agora é escolher o nicho inicial. A recomendação é começar por e-commerce, especialmente moda, beleza, perfumaria ou varejo especializado, porque os dados de compra, recompra, carrinho, produto e campanha são objetivos e permitem provar ROI mais rápido.",
    )

    footer = doc.sections[0].footer.paragraphs[0]
    footer.alignment = WD_ALIGN_PARAGRAPH.CENTER
    r = footer.add_run("PersonaPulse AI - Documento de construção do MVP")
    style_run(r, color="777777", size=9)

    doc.save(OUTPUT)


if __name__ == "__main__":
    build()

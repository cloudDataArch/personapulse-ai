from copy import deepcopy

from docx import Document
from docx.enum.table import WD_CELL_VERTICAL_ALIGNMENT, WD_TABLE_ALIGNMENT
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from docx.shared import Inches, Pt, RGBColor
from pathlib import Path


ROOT = Path(__file__).resolve().parents[1]
INPUT = ROOT / "DOCUMENTOS" / "PersonaPulse_AI_Organograma_Etapas_MVP.docx"
OUTPUT = ROOT / "DOCUMENTOS" / "PersonaPulse_AI_Organograma_Etapas_MVP_atualizado.docx"

BLUE = "2E74B5"
DARK_BLUE = "1F4D78"
INK = "0B2545"
LIGHT_BLUE = "E8EEF5"
LIGHT_GRAY = "F2F4F7"
PALE = "F7F9FC"
GREEN = "EAF4EE"


def style_run(run, bold=False, color=None, size=None):
    run.bold = bold
    if color:
        run.font.color.rgb = RGBColor.from_string(color)
    if size:
        run.font.size = Pt(size)


def set_cell_shading(cell, fill):
    tc_pr = cell._tc.get_or_add_tcPr()
    shd = tc_pr.find(qn("w:shd"))
    if shd is None:
        shd = OxmlElement("w:shd")
        tc_pr.append(shd)
    shd.set(qn("w:fill"), fill)


def set_cell_margins(cell, top=100, start=120, bottom=100, end=120):
    tc_pr = cell._tc.get_or_add_tcPr()
    tc_mar = tc_pr.first_child_found_in("w:tcMar")
    if tc_mar is None:
        tc_mar = OxmlElement("w:tcMar")
        tc_pr.append(tc_mar)
    for m, v in [("top", top), ("start", start), ("bottom", bottom), ("end", end)]:
        node = tc_mar.find(qn(f"w:{m}"))
        if node is None:
            node = OxmlElement(f"w:{m}")
            tc_mar.append(node)
        node.set(qn("w:w"), str(v))
        node.set(qn("w:type"), "dxa")


def set_table_width(table, widths):
    table.alignment = WD_TABLE_ALIGNMENT.CENTER
    table.autofit = False
    for row in table.rows:
        for idx, width in enumerate(widths):
            cell = row.cells[idx]
            cell.width = Inches(width)
            tc_pr = cell._tc.get_or_add_tcPr()
            tc_w = tc_pr.find(qn("w:tcW"))
            if tc_w is None:
                tc_w = OxmlElement("w:tcW")
                tc_pr.append(tc_w)
            tc_w.set(qn("w:w"), str(int(width * 1440)))
            tc_w.set(qn("w:type"), "dxa")


def insert_after(paragraph, element):
    paragraph._p.addnext(element)


def make_paragraph_after(paragraph, text="", style=None):
    new_p = OxmlElement("w:p")
    insert_after(paragraph, new_p)
    p = paragraph._parent.add_paragraph()
    p._p.getparent().remove(p._p)
    new_p.addnext(p._p)
    paragraph._p.getparent().remove(new_p)
    if style:
        p.style = style
    if text:
        p.add_run(text)
    return p


def add_table_after(paragraph, rows, cols, widths):
    doc_part = paragraph._parent
    table = doc_part.add_table(rows=rows, cols=cols, width=Inches(6.5))
    table.style = "Table Grid"
    set_table_width(table, widths)
    tbl = table._tbl
    tbl.getparent().remove(tbl)
    paragraph._p.addnext(tbl)
    return table


def format_box(cell, title, body, fill):
    set_cell_shading(cell, fill)
    set_cell_margins(cell, top=130, bottom=130, start=150, end=150)
    cell.vertical_alignment = WD_CELL_VERTICAL_ALIGNMENT.CENTER
    p = cell.paragraphs[0]
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    r = p.add_run(title)
    style_run(r, bold=True, color=INK, size=10.5)
    p.add_run("\n")
    r = p.add_run(body)
    style_run(r, color="333333", size=9)


def build_orgchart(doc):
    target = None
    for p in doc.paragraphs:
        if p.text.strip() == "Organograma Operacional Simplificado":
            target = p
            break
    if target is None:
        raise RuntimeError("Heading not found")

    intro = make_paragraph_after(
        target,
        "O desenho abaixo mostra a estrutura de construção do projeto, separando decisão, produto, tecnologia, experiência, compliance e execução técnica.",
    )

    table = add_table_after(intro, 5, 4, [1.625, 1.625, 1.625, 1.625])
    for row in table.rows:
        for cell in row.cells:
            cell.text = ""

    # Row 0: sponsor merged.
    sponsor = table.cell(0, 0).merge(table.cell(0, 3))
    format_box(
        sponsor,
        "Fundador / Sponsor",
        "Visao, orcamento, decisao final e validacao comercial",
        LIGHT_BLUE,
    )

    # Row 1: direct leadership.
    direct = [
        ("Product Manager", "Escopo, roadmap, backlog e aceite", LIGHT_GRAY),
        ("Tech Lead", "Arquitetura e lideranca tecnica", LIGHT_GRAY),
        ("UX/UI Designer", "Prototipo, telas e experiencia", LIGHT_GRAY),
        ("LGPD / Juridico", "Privacidade, consentimento e compliance", LIGHT_GRAY),
    ]
    for idx, item in enumerate(direct):
        format_box(table.cell(1, idx), *item)

    # Row 2: technical execution under Tech Lead.
    exec_row = table.cell(2, 0).merge(table.cell(2, 3))
    format_box(
        exec_row,
        "Execucao Tecnica sob o Tech Lead",
        "Frontend, Backend, Dados e IA, DevOps / Cloud e QA / Testes",
        PALE,
    )

    technical = [
        ("Frontend", "Dashboard, segmentos e campanhas", "FFFFFF"),
        ("Backend", "APIs, banco, autenticacao e integracoes", "FFFFFF"),
        ("Dados e IA", "Segmentacao, scoring, prompts e recomendacoes", "FFFFFF"),
        ("DevOps / QA", "Deploy, logs, backups, testes e qualidade", "FFFFFF"),
    ]
    for idx, item in enumerate(technical):
        format_box(table.cell(3, idx), *item)

    outputs = table.cell(4, 0).merge(table.cell(4, 3))
    format_box(
        outputs,
        "Entregaveis do MVP",
        "Painel web, motor de IA generativa, modulo LGPD, recomendacoes diarias e piloto com clientes",
        GREEN,
    )

    # Add textual hierarchy after the table for readers who prefer the structure in outline form.
    after = table._tbl
    spacer = OxmlElement("w:p")
    after.addnext(spacer)
    p = doc.add_paragraph()
    p._p.getparent().remove(p._p)
    spacer.addnext(p._p)
    spacer.getparent().remove(spacer)
    p.style = "Heading 3"
    p.add_run("Hierarquia do Projeto")

    hierarchy = [
        "Fundador / Sponsor",
        "Product Manager: escopo do MVP, roadmap e validacao com clientes.",
        "UX/UI Designer: prototipo navegavel, telas do painel e design system.",
        "Tech Lead: arquitetura, decisao tecnica e lideranca de implementacao.",
        "Frontend: dashboard, segmentos, campanhas e interface SaaS.",
        "Backend: APIs, banco, autenticacao, ingestao e integracoes.",
        "Dados e IA: segmentacao, scoring, prompts e IA generativa.",
        "DevOps / Cloud: deploy, monitoramento, logs e backups.",
        "QA / Testes: testes, bugs, regressao e validacao.",
        "LGPD / Juridico: consentimento, politica de privacidade, exclusao de dados e auditoria.",
    ]
    previous = p
    for item in hierarchy:
        new_p = make_paragraph_after(previous, item, style="List Bullet")
        previous = new_p


def build():
    doc = Document(INPUT)
    build_orgchart(doc)
    doc.save(OUTPUT)


if __name__ == "__main__":
    build()

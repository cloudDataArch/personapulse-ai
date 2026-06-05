from textwrap import wrap

from PIL import Image, ImageDraw, ImageFont
from docx import Document
from docx.shared import Inches, Pt


BASE_DOC = r"C:\Users\Celio\Documents\Codex\2026-05-31\eu-quero-ajuda-para-criar-uma\outputs\PersonaPulse_AI_Organograma_Etapas_MVP.docx"
OUTPUT_DOC = r"C:\Users\Celio\Documents\Codex\2026-05-31\eu-quero-ajuda-para-criar-uma\outputs\PersonaPulse_AI_Organograma_Etapas_MVP_atualizado_v2.docx"
ORG_IMAGE = r"C:\Users\Celio\Documents\Codex\2026-05-31\eu-quero-ajuda-para-criar-uma\work\organograma_personapulse.png"


INK = "#0B2545"
BLUE = "#2E74B5"
DARK_BLUE = "#1F4D78"
LIGHT_BLUE = "#E8EEF5"
LIGHT_GRAY = "#F2F4F7"
PALE = "#F7F9FC"
GREEN = "#EAF4EE"
BORDER = "#8BA6C1"


def font(size, bold=False):
    path = r"C:\Windows\Fonts\arialbd.ttf" if bold else r"C:\Windows\Fonts\arial.ttf"
    return ImageFont.truetype(path, size)


def multiline_center(draw, box, title, body, fill, outline=BORDER):
    x1, y1, x2, y2 = box
    draw.rounded_rectangle(box, radius=18, fill=fill, outline=outline, width=3)
    title_font = font(28, bold=True)
    body_font = font(22)
    lines = [title]
    max_chars = max(16, int((x2 - x1) / 15))
    lines.extend(wrap(body, width=max_chars))
    line_heights = []
    for i, line in enumerate(lines):
        f = title_font if i == 0 else body_font
        bbox = draw.textbbox((0, 0), line, font=f)
        line_heights.append(bbox[3] - bbox[1])
    total_h = sum(line_heights) + (len(lines) - 1) * 8
    y = y1 + ((y2 - y1) - total_h) / 2
    for i, line in enumerate(lines):
        f = title_font if i == 0 else body_font
        fill_color = INK if i == 0 else "#333333"
        bbox = draw.textbbox((0, 0), line, font=f)
        w = bbox[2] - bbox[0]
        draw.text((x1 + ((x2 - x1) - w) / 2, y), line, font=f, fill=fill_color)
        y += line_heights[i] + 8


def arrow(draw, start, end):
    draw.line([start, end], fill=DARK_BLUE, width=4)
    ex, ey = end
    sx, sy = start
    if ey >= sy:
        pts = [(ex, ey), (ex - 10, ey - 18), (ex + 10, ey - 18)]
    else:
        pts = [(ex, ey), (ex - 10, ey + 18), (ex + 10, ey + 18)]
    draw.polygon(pts, fill=DARK_BLUE)


def make_org_image():
    img = Image.new("RGB", (1800, 1280), "white")
    draw = ImageDraw.Draw(img)

    title_font = font(44, bold=True)
    subtitle_font = font(25)
    draw.text((70, 45), "Organograma do Projeto PersonaPulse AI", font=title_font, fill=INK)
    draw.text(
        (70, 100),
        "Estrutura de decisao, produto, tecnologia, experiencia, compliance e execucao tecnica",
        font=subtitle_font,
        fill="#555555",
    )

    sponsor = (550, 170, 1250, 290)
    multiline_center(draw, sponsor, "Fundador / Sponsor", "Visao, orcamento e decisao final", LIGHT_BLUE)

    row_y1, row_y2 = 400, 540
    boxes = [
        ((80, row_y1, 430, row_y2), "Product Manager", "Escopo, roadmap, backlog e aceite"),
        ((520, row_y1, 870, row_y2), "Tech Lead", "Arquitetura e lideranca tecnica"),
        ((960, row_y1, 1310, row_y2), "UX/UI Designer", "Prototipo, telas e experiencia"),
        ((1400, row_y1, 1750, row_y2), "LGPD / Juridico", "Privacidade, consentimento e compliance"),
    ]
    for box, title, body in boxes:
        multiline_center(draw, box, title, body, LIGHT_GRAY)
        arrow(draw, ((sponsor[0] + sponsor[2]) // 2, sponsor[3]), ((box[0] + box[2]) // 2, box[1]))

    execution = (390, 650, 1410, 770)
    multiline_center(
        draw,
        execution,
        "Execucao Tecnica sob o Tech Lead",
        "Frontend, Backend, Dados e IA, DevOps / Cloud e QA / Testes",
        PALE,
    )
    arrow(draw, (695, row_y2), (900, execution[1]))

    tech_y1, tech_y2 = 870, 1030
    tech_boxes = [
        ((80, tech_y1, 430, tech_y2), "Frontend", "Dashboard, segmentos e campanhas"),
        ((520, tech_y1, 870, tech_y2), "Backend", "APIs, banco, autenticacao e integracoes"),
        ((960, tech_y1, 1310, tech_y2), "Dados e IA", "Segmentacao, scoring, prompts e recomendacoes"),
        ((1400, tech_y1, 1750, tech_y2), "DevOps / QA", "Deploy, logs, backups, testes e qualidade"),
    ]
    for box, title, body in tech_boxes:
        multiline_center(draw, box, title, body, "white")
        arrow(draw, ((execution[0] + execution[2]) // 2, execution[3]), ((box[0] + box[2]) // 2, box[1]))

    deliverables = (340, 1130, 1460, 1235)
    multiline_center(
        draw,
        deliverables,
        "Entregaveis do MVP",
        "Painel web, motor de IA generativa, modulo LGPD, recomendacoes diarias e piloto com clientes",
        GREEN,
    )
    for box, _, _ in tech_boxes:
        arrow(draw, ((box[0] + box[2]) // 2, box[3]), ((deliverables[0] + deliverables[2]) // 2, deliverables[1]))

    img.save(ORG_IMAGE, quality=95)


def insert_after(ref_paragraph, new_paragraph):
    ref_paragraph._p.addnext(new_paragraph._p)
    return new_paragraph


def add_paragraph_after(doc, ref, text="", style=None):
    p = doc.add_paragraph()
    if style:
        p.style = style
    if text:
        p.add_run(text)
    return insert_after(ref, p)


def add_picture_after(doc, ref, path):
    p = doc.add_paragraph()
    p.alignment = 1
    p.add_run().add_picture(path, width=Inches(6.5))
    return insert_after(ref, p)


def remove_original_simple_hierarchy(doc):
    start_idx = None
    end_idx = None
    paragraphs = list(doc.paragraphs)
    for i, p in enumerate(paragraphs):
        if p.text.strip() == "Organograma Operacional Simplificado":
            start_idx = i
        elif start_idx is not None and p.style.name == "Heading 1" and p.text.strip().startswith("2. "):
            end_idx = i
            break
    if start_idx is None or end_idx is None:
        return
    # Remove paragraphs between this heading and the next section; the new drawing and hierarchy replace them.
    for p in paragraphs[start_idx + 1 : end_idx]:
        parent = p._p.getparent()
        if parent is not None:
            parent.remove(p._p)


def build_doc():
    make_org_image()
    doc = Document(BASE_DOC)
    remove_original_simple_hierarchy(doc)

    heading = None
    for p in doc.paragraphs:
        if p.text.strip() == "Organograma Operacional Simplificado":
            heading = p
            break
    if heading is None:
        raise RuntimeError("Heading not found")

    current = add_paragraph_after(
        doc,
        heading,
        "Primeiro, o organograma visual do projeto. Em seguida, a mesma estrutura aparece em formato hierarquico para leitura operacional.",
    )
    current = add_picture_after(doc, current, ORG_IMAGE)
    current = add_paragraph_after(doc, current, "Formato Hierarquico", style="Heading 3")

    hierarchy = [
        "Fundador / Sponsor",
        "Product Manager: escopo do MVP, roadmap e validacao com clientes.",
        "UX/UI Designer: prototipo navegavel, telas do painel e design system.",
        "Tech Lead: arquitetura, decisao tecnica e lideranca de implementacao.",
        "Frontend: dashboard, segmentos, campanhas e interface SaaS.",
        "Backend: APIs, banco, autenticacao, ingestao e integracoes.",
        "Dados e IA: segmentacao, scoring, prompts e IA generativa.",
        "DevOps / Cloud: deploy, monitoramento, logs e backups.",
        "QA / Testes: testes, bugs, regressao e validacao.",
        "LGPD / Juridico: consentimento, politica de privacidade, exclusao de dados e auditoria.",
    ]
    for item in hierarchy:
        current = add_paragraph_after(doc, current, item, style="List Bullet")

    doc.save(OUTPUT_DOC)


if __name__ == "__main__":
    build_doc()

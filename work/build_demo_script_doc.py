from docx import Document
from docx.enum.table import WD_CELL_VERTICAL_ALIGNMENT, WD_TABLE_ALIGNMENT
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from docx.shared import Inches, Pt, RGBColor


OUTPUT = r"C:\Users\Celio\Documents\Codex\2026-05-31\eu-quero-ajuda-para-criar-uma\outputs\PersonaPulse_AI_Roteiro_Demo_Produto.docx"

BLUE = "2E74B5"
DARK_BLUE = "1F4D78"
INK = "0B2545"
LIGHT_BLUE = "E8EEF5"
LIGHT_GRAY = "F2F4F7"
PALE = "F7F9FC"
GREEN = "EAF4EE"
GOLD = "FFF7DC"


def style_run(run, bold=False, color=None, size=None):
    run.bold = bold
    if color:
        run.font.color.rgb = RGBColor.from_string(color)
    if size:
        run.font.size = Pt(size)


def set_cell_shading(cell, fill):
    tc_pr = cell._tc.get_or_add_tcPr()
    shd = tc_pr.find(qn("w:shd"))
    if shd is None:
        shd = OxmlElement("w:shd")
        tc_pr.append(shd)
    shd.set(qn("w:fill"), fill)


def set_cell_margins(cell, top=100, start=120, bottom=100, end=120):
    tc_pr = cell._tc.get_or_add_tcPr()
    tc_mar = tc_pr.first_child_found_in("w:tcMar")
    if tc_mar is None:
        tc_mar = OxmlElement("w:tcMar")
        tc_pr.append(tc_mar)
    for m, v in [("top", top), ("start", start), ("bottom", bottom), ("end", end)]:
        node = tc_mar.find(qn(f"w:{m}"))
        if node is None:
            node = OxmlElement(f"w:{m}")
            tc_mar.append(node)
        node.set(qn("w:w"), str(v))
        node.set(qn("w:type"), "dxa")


def set_table_width(table, widths):
    table.alignment = WD_TABLE_ALIGNMENT.CENTER
    table.autofit = False
    for row in table.rows:
        for idx, width in enumerate(widths):
            cell = row.cells[idx]
            cell.width = Inches(width)
            tc_pr = cell._tc.get_or_add_tcPr()
            tc_w = tc_pr.find(qn("w:tcW"))
            if tc_w is None:
                tc_w = OxmlElement("w:tcW")
                tc_pr.append(tc_w)
            tc_w.set(qn("w:w"), str(int(width * 1440)))
            tc_w.set(qn("w:type"), "dxa")


def add_heading(doc, text, level=1):
    p = doc.add_paragraph(style=f"Heading {level}")
    p.add_run(text)
    return p


def add_body(doc, text):
    p = doc.add_paragraph()
    p.add_run(text)
    return p


def add_bullets(doc, items):
    for item in items:
        p = doc.add_paragraph(style="List Bullet")
        p.add_run(item)


def add_callout(doc, title, body, fill=PALE):
    table = doc.add_table(rows=1, cols=1)
    table.style = "Table Grid"
    set_table_width(table, [6.5])
    cell = table.cell(0, 0)
    set_cell_shading(cell, fill)
    set_cell_margins(cell, top=150, bottom=150, start=180, end=180)
    p = cell.paragraphs[0]
    r = p.add_run(title)
    style_run(r, bold=True, color=INK, size=11)
    p.add_run("\n" + body)
    doc.add_paragraph()


def add_table(doc, headers, rows, widths):
    table = doc.add_table(rows=1, cols=len(headers))
    table.style = "Table Grid"
    set_table_width(table, widths)
    for i, h in enumerate(headers):
        cell = table.rows[0].cells[i]
        set_cell_shading(cell, LIGHT_GRAY)
        set_cell_margins(cell)
        cell.vertical_alignment = WD_CELL_VERTICAL_ALIGNMENT.CENTER
        r = cell.paragraphs[0].add_run(h)
        style_run(r, bold=True, color=INK, size=9.5)
    for row in rows:
        cells = table.add_row().cells
        for i, value in enumerate(row):
            set_cell_margins(cells[i])
            cells[i].vertical_alignment = WD_CELL_VERTICAL_ALIGNMENT.CENTER
            r = cells[i].paragraphs[0].add_run(value)
            style_run(r, size=9.2)
    doc.add_paragraph()
    return table


def build():
    doc = Document()
    sec = doc.sections[0]
    sec.page_width = Inches(8.5)
    sec.page_height = Inches(11)
    sec.top_margin = Inches(1)
    sec.bottom_margin = Inches(1)
    sec.left_margin = Inches(1)
    sec.right_margin = Inches(1)

    styles = doc.styles
    normal = styles["Normal"]
    normal.font.name = "Calibri"
    normal.font.size = Pt(11)
    normal.paragraph_format.space_after = Pt(6)
    normal.paragraph_format.line_spacing = 1.10

    for name, size, color, before, after in [
        ("Heading 1", 16, BLUE, 16, 8),
        ("Heading 2", 13, BLUE, 12, 6),
        ("Heading 3", 12, DARK_BLUE, 8, 4),
    ]:
        st = styles[name]
        st.font.name = "Calibri"
        st.font.size = Pt(size)
        st.font.color.rgb = RGBColor.from_string(color)
        st.font.bold = True
        st.paragraph_format.space_before = Pt(before)
        st.paragraph_format.space_after = Pt(after)

    title = doc.add_paragraph()
    title.alignment = WD_ALIGN_PARAGRAPH.CENTER
    r = title.add_run("PersonaPulse AI")
    style_run(r, bold=True, color=INK, size=24)
    subtitle = doc.add_paragraph()
    subtitle.alignment = WD_ALIGN_PARAGRAPH.CENTER
    r = subtitle.add_run("Roteiro de demonstracao do produto")
    style_run(r, color=DARK_BLUE, size=13)

    add_callout(
        doc,
        "Objetivo da demo",
        "Mostrar, em 8 a 12 minutos, como uma empresa importa dados proprios, entende seus clientes, recebe recomendacoes da IA e gera campanhas prontas para vender mais com governanca LGPD.",
        fill=LIGHT_BLUE,
    )

    add_heading(doc, "1. Cenario da Demonstracao", 1)
    add_body(doc, "Use um exemplo simples e comercialmente facil de entender: uma loja de perfumes premium que vende pelo site, WhatsApp e campanhas de anuncios.")
    add_bullets(doc, [
        "Empresa ficticia: Essenza Prime.",
        "Problema: muitos dados de clientes, mas campanhas genericas e pouco aproveitamento de recompra.",
        "Base simulada: clientes, produtos, compras, carrinhos abandonados e consentimentos.",
        "Promessa da demo: transformar dados consentidos em segmentos, recomendacoes e campanhas prontas.",
    ])

    add_heading(doc, "2. Abertura da Apresentacao", 1)
    add_callout(
        doc,
        "Fala sugerida",
        "Hoje as empresas gastam muito com anuncios, mas ainda tomam decisoes com dados fragmentados. O PersonaPulse AI junta dados proprios e consentidos da empresa, identifica oportunidades de venda e gera campanhas automaticamente com IA generativa.",
        fill=GREEN,
    )
    add_body(doc, "Tempo recomendado: 45 a 60 segundos.")

    add_heading(doc, "3. Roteiro de Cliques e Fala", 1)
    rows = [
        ("1", "Dashboard", "Abrir a tela inicial.", "Aqui vemos a saude comercial do dia: clientes analisados, receita influenciada, campanhas prontas e conformidade LGPD.", "Mostra valor executivo logo no inicio."),
        ("2", "Dashboard", "Clicar em Simular CSV.", "Agora simulo a importacao de uma base atualizada. O sistema processa os novos clientes e atualiza os indicadores.", "Demonstra ingestao de dados sem entrar em detalhe tecnico."),
        ("3", "Segmentos", "Abrir o menu Segmentos.", "A IA organiza a base em perfis comportamentais: clientes premium, carrinho abandonado, recompra provavel, sensiveis a desconto e inativos.", "Mostra que o produto transforma dados em grupos acionaveis."),
        ("4", "Segmentos", "Apontar a coluna Acao sugerida.", "Cada segmento ja vem com uma recomendacao de acao. O usuario nao precisa interpretar relatorios do zero.", "Diferencia de analytics tradicional."),
        ("5", "Campanhas", "Abrir Campanhas.", "Agora escolhemos o segmento, o canal, o produto e o tom da mensagem.", "Prepara a geracao com IA."),
        ("6", "Campanhas", "Clicar em Gerar campanha.", "A IA cria uma campanha pronta para WhatsApp, e-mail, Ads ou SEO, com titulo, mensagem, chamada para acao e justificativa.", "Demonstra IA generativa aplicada a venda."),
        ("7", "Recomendacoes", "Abrir Recomendacoes.", "Aqui a plataforma vira um copiloto de marketing: mostra o que fazer primeiro, por que fazer e qual teste rodar.", "Mostra decisao diaria, nao apenas relatorio."),
        ("8", "LGPD", "Abrir LGPD.", "Toda recomendacao parte de dados com origem, finalidade e consentimento. A empresa consegue controlar opt-out, retencao e auditoria.", "Reduz preocupacao legal e aumenta confianca."),
        ("9", "Fechamento", "Voltar ao Dashboard.", "O resultado e uma plataforma que conecta dados, IA e acao comercial em uma rotina diaria simples.", "Fecha a narrativa do produto."),
    ]
    add_table(doc, ["Passo", "Tela", "Acao", "Fala sugerida", "Mensagem que fica"], rows, [0.55, 1.00, 1.35, 2.40, 1.20])

    add_heading(doc, "4. Historia que a Demo Deve Contar", 1)
    add_body(doc, "A demonstracao precisa parecer uma jornada, nao uma lista de telas. A historia recomendada e:")
    add_bullets(doc, [
        "A empresa tem dados, mas nao sabe exatamente que campanha fazer hoje.",
        "Ela importa uma base de clientes e compras.",
        "O PersonaPulse AI identifica segmentos e oportunidades.",
        "A IA recomenda a melhor acao por prioridade.",
        "A IA gera campanhas prontas para cada canal.",
        "A empresa opera tudo com consentimento, auditoria e LGPD.",
    ])

    add_heading(doc, "5. Perguntas que Podem Surgir", 1)
    faq_rows = [
        ("De onde vem os dados?", "De fontes proprias e autorizadas: CSV, e-commerce, CRM, historico de compras, eventos do site e consentimentos."),
        ("Isso espiona o usuario?", "Nao. O posicionamento correto e inteligencia comportamental consentida, baseada em dados proprios da empresa."),
        ("A IA decide sozinha?", "Nao no MVP. Ela recomenda e gera rascunhos, mas o humano aprova antes de publicar."),
        ("Qual o primeiro nicho?", "E-commerce de moda, beleza, perfumaria ou varejo especializado, porque o ROI e mais facil de medir."),
        ("Qual o diferencial?", "Unir segmentacao, recomendacao diaria, IA generativa e governanca LGPD em uma unica experiencia."),
    ]
    add_table(doc, ["Pergunta", "Resposta recomendada"], faq_rows, [2.00, 4.50])

    add_heading(doc, "6. Fechamento Comercial", 1)
    add_callout(
        doc,
        "Fala de encerramento",
        "O PersonaPulse AI nao e apenas um painel. Ele e um copiloto de crescimento: entende a base, encontra oportunidades, recomenda a proxima acao e gera campanhas prontas, sempre com privacidade e consentimento.",
        fill=GOLD,
    )

    add_heading(doc, "7. Proximo Passo Depois da Demo", 1)
    add_bullets(doc, [
        "Escolher um nicho inicial para validar.",
        "Criar uma base CSV realista de exemplo.",
        "Adaptar textos do prototipo para esse nicho.",
        "Montar uma apresentacao curta para potenciais clientes ou socios.",
        "Depois conectar upload CSV real e banco de dados.",
    ])

    footer = doc.sections[0].footer.paragraphs[0]
    footer.alignment = WD_ALIGN_PARAGRAPH.CENTER
    r = footer.add_run("PersonaPulse AI - Roteiro de demo do produto")
    style_run(r, color="777777", size=9)
    doc.save(OUTPUT)


if __name__ == "__main__":
    build()

from docx import Document
from docx.enum.table import WD_CELL_VERTICAL_ALIGNMENT, WD_TABLE_ALIGNMENT
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from docx.shared import Inches, Pt, RGBColor


OUTPUT = r"C:\Users\Celio\Documents\Codex\2026-05-31\eu-quero-ajuda-para-criar-uma\outputs\PersonaPulse_AI_Comparativo_Ferramentas_Atuais.docx"

BLUE = "2E74B5"
DARK_BLUE = "1F4D78"
INK = "0B2545"
LIGHT_BLUE = "E8EEF5"
LIGHT_GRAY = "F2F4F7"
PALE = "F7F9FC"
GREEN = "EAF4EE"


def style_run(run, bold=False, color=None, size=None):
    run.bold = bold
    if color:
        run.font.color.rgb = RGBColor.from_string(color)
    if size:
        run.font.size = Pt(size)


def set_cell_shading(cell, fill):
    tc_pr = cell._tc.get_or_add_tcPr()
    shd = tc_pr.find(qn("w:shd"))
    if shd is None:
        shd = OxmlElement("w:shd")
        tc_pr.append(shd)
    shd.set(qn("w:fill"), fill)


def set_cell_margins(cell, top=100, start=120, bottom=100, end=120):
    tc_pr = cell._tc.get_or_add_tcPr()
    tc_mar = tc_pr.first_child_found_in("w:tcMar")
    if tc_mar is None:
        tc_mar = OxmlElement("w:tcMar")
        tc_pr.append(tc_mar)
    for m, v in [("top", top), ("start", start), ("bottom", bottom), ("end", end)]:
        node = tc_mar.find(qn(f"w:{m}"))
        if node is None:
            node = OxmlElement(f"w:{m}")
            tc_mar.append(node)
        node.set(qn("w:w"), str(v))
        node.set(qn("w:type"), "dxa")


def set_table_width(table, widths):
    table.alignment = WD_TABLE_ALIGNMENT.CENTER
    table.autofit = False
    for row in table.rows:
        for idx, width in enumerate(widths):
            cell = row.cells[idx]
            cell.width = Inches(width)
            tc_pr = cell._tc.get_or_add_tcPr()
            tc_w = tc_pr.find(qn("w:tcW"))
            if tc_w is None:
                tc_w = OxmlElement("w:tcW")
                tc_pr.append(tc_w)
            tc_w.set(qn("w:w"), str(int(width * 1440)))
            tc_w.set(qn("w:type"), "dxa")


def add_heading(doc, text, level=1):
    p = doc.add_paragraph(style=f"Heading {level}")
    p.add_run(text)
    return p


def add_body(doc, text):
    p = doc.add_paragraph()
    p.add_run(text)
    return p


def add_bullets(doc, items):
    for item in items:
        p = doc.add_paragraph(style="List Bullet")
        p.add_run(item)


def add_callout(doc, title, body, fill=PALE):
    table = doc.add_table(rows=1, cols=1)
    table.style = "Table Grid"
    set_table_width(table, [6.5])
    cell = table.cell(0, 0)
    set_cell_shading(cell, fill)
    set_cell_margins(cell, top=150, bottom=150, start=180, end=180)
    p = cell.paragraphs[0]
    r = p.add_run(title)
    style_run(r, bold=True, color=INK, size=11)
    p.add_run("\n" + body)
    doc.add_paragraph()


def add_table(doc, headers, rows, widths):
    table = doc.add_table(rows=1, cols=len(headers))
    table.style = "Table Grid"
    set_table_width(table, widths)
    for i, h in enumerate(headers):
        cell = table.rows[0].cells[i]
        set_cell_shading(cell, LIGHT_GRAY)
        set_cell_margins(cell)
        cell.vertical_alignment = WD_CELL_VERTICAL_ALIGNMENT.CENTER
        r = cell.paragraphs[0].add_run(h)
        style_run(r, bold=True, color=INK, size=9.5)
    for row in rows:
        cells = table.add_row().cells
        for i, value in enumerate(row):
            set_cell_margins(cells[i])
            cells[i].vertical_alignment = WD_CELL_VERTICAL_ALIGNMENT.CENTER
            r = cells[i].paragraphs[0].add_run(value)
            style_run(r, size=9.2)
    doc.add_paragraph()
    return table


def build():
    doc = Document()
    sec = doc.sections[0]
    sec.page_width = Inches(8.5)
    sec.page_height = Inches(11)
    sec.top_margin = Inches(1)
    sec.bottom_margin = Inches(1)
    sec.left_margin = Inches(1)
    sec.right_margin = Inches(1)

    styles = doc.styles
    normal = styles["Normal"]
    normal.font.name = "Calibri"
    normal.font.size = Pt(11)
    normal.paragraph_format.space_after = Pt(6)
    normal.paragraph_format.line_spacing = 1.10
    for name, size, color, before, after in [
        ("Heading 1", 16, BLUE, 16, 8),
        ("Heading 2", 13, BLUE, 12, 6),
        ("Heading 3", 12, DARK_BLUE, 8, 4),
    ]:
        st = styles[name]
        st.font.name = "Calibri"
        st.font.size = Pt(size)
        st.font.color.rgb = RGBColor.from_string(color)
        st.font.bold = True
        st.paragraph_format.space_before = Pt(before)
        st.paragraph_format.space_after = Pt(after)

    title = doc.add_paragraph()
    title.alignment = WD_ALIGN_PARAGRAPH.CENTER
    r = title.add_run("PersonaPulse AI")
    style_run(r, bold=True, color=INK, size=24)
    subtitle = doc.add_paragraph()
    subtitle.alignment = WD_ALIGN_PARAGRAPH.CENTER
    r = subtitle.add_run("Comparativo com ferramentas atuais de navegadores, redes sociais e dispositivos")
    style_run(r, color=DARK_BLUE, size=13)

    add_callout(
        doc,
        "Diferenca central",
        "As ferramentas atuais observam partes do comportamento do usuario. A PersonaPulse AI deve ser posicionada como uma camada unificada de inteligencia comportamental consentida, capaz de juntar dados proprios, interpretar perfis e gerar acoes de marketing automaticamente.",
        fill=LIGHT_BLUE,
    )

    add_heading(doc, "1. Comparativo Geral", 1)
    rows = [
        ("Google Analytics", "Mede eventos em sites e apps, como paginas, cliques, sessoes e conversoes.", "Fica muito focado no comportamento dentro do site/app.", "Junta dados comerciais, CRM, compras, campanhas e preferencias para gerar recomendacoes acionaveis."),
        ("Meta Pixel / pixels de ads", "Rastreia acoes no site, como visita, carrinho e compra, para otimizar anuncios.", "Depende do ambiente da plataforma e de eventos configurados.", "Nao e apenas pixel; funciona como motor de analise, segmentacao e geracao de campanhas."),
        ("Redes sociais", "Conhecem curtidas, tempo de video, interacoes, interesses e engajamento dentro da propria plataforma.", "Cada rede enxerga principalmente seu proprio ecossistema.", "Pode cruzar dados proprios da empresa, como CRM, historico de compra e campanhas, sem depender so de uma rede."),
        ("Navegadores / cookies", "Usam cookies, permissoes e mecanismos de privacidade para controlar rastreamento entre sites.", "O rastreamento esta mais limitado por privacidade, bloqueios e consentimento.", "Deve ser baseada em dados first-party, ou seja, dados da propria empresa com autorizacao."),
        ("Dispositivos moveis", "Apps podem coletar dados, mas sistemas como iOS exigem permissao para rastrear entre apps e sites.", "Forte dependencia de permissao do usuario e regras da loja.", "Nao depende de coleta oculta no dispositivo; usa consentimento claro e dados declarados/comerciais."),
        ("CRMs", "Guardam dados de clientes, leads, historico comercial e atendimento.", "Muitas vezes nao possuem IA generativa ou analise comportamental avancada.", "Usa esses dados para sugerir campanhas, mensagens, ofertas e segmentos automaticamente."),
        ("Automacao de marketing", "Dispara e-mails, fluxos e segmentacoes.", "Normalmente depende de configuracao manual e regras fixas.", "Atua de forma consultiva: recomenda o que fazer hoje, para qual publico, com qual mensagem."),
    ]
    add_table(doc, ["Ferramenta atual", "O que faz hoje", "Limitacao", "Diferenca da PersonaPulse AI"], rows, [1.25, 1.80, 1.55, 1.90])

    add_heading(doc, "2. Diferenca Estrategica", 1)
    add_body(doc, "As ferramentas atuais costumam ser:")
    add_bullets(doc, [
        "Fragmentadas: cada uma ve um pedaco do usuario.",
        "Reativas: mostram relatorios depois que algo aconteceu.",
        "Dependentes de configuracao manual.",
        "Limitadas ao proprio canal.",
        "Mais focadas em mensuracao do que em decisao.",
    ])
    add_body(doc, "A PersonaPulse AI deve ser:")
    add_bullets(doc, [
        "Unificada: junta dados de site, CRM, compras, campanhas e atendimento.",
        "Preditiva: identifica chance de compra, recompra, abandono e interesse.",
        "Generativa: cria textos, anuncios, e-mails, WhatsApp e SEO.",
        "Acionavel: recomenda o que fazer todos os dias.",
        "Orientada a LGPD: com consentimento, opt-out, origem dos dados e exclusao.",
    ])

    add_heading(doc, "3. Ponto de Atencao Legal e Reputacional", 1)
    add_body(
        doc,
        "A ideia original de observar o usuario em muitos ambientes, como navegador, redes, Wi-Fi publico, IoT, mapas, cameras e comportamento fisico, tem alto risco legal, reputacional e operacional. Mesmo quando dados sao mascarados, ainda pode haver identificacao indireta e tratamento de dados sensiveis.",
    )
    add_callout(
        doc,
        "Posicionamento recomendado",
        "Nao somos uma ferramenta de vigilancia. Somos uma plataforma de inteligencia comportamental consentida, baseada em dados proprios da empresa, para transformar informacao em campanhas e recomendacoes de venda.",
        fill=GREEN,
    )

    add_heading(doc, "4. Conclusao para o Projeto", 1)
    add_body(
        doc,
        "A vantagem competitiva da PersonaPulse AI nao deve estar em coletar mais dados do que todos, mas em transformar dados permitidos em decisoes melhores. O produto precisa se diferenciar por integracao, inteligencia, geracao automatica de campanhas, facilidade de uso e governanca de privacidade.",
    )

    add_heading(doc, "5. Fontes Consultadas", 1)
    add_bullets(doc, [
        "Google Analytics - como funciona: https://support.google.com/analytics/answer/12159447?hl=en",
        "Google Analytics - eventos: https://support.google.com/analytics/answer/9322688?hl=en",
        "Apple App Tracking Transparency: https://support.apple.com/en-euro/102420",
        "Apple Developer - User Privacy and Data Use: https://developer.apple.com/app-store/user-privacy-and-data-use/",
        "Google Privacy Sandbox: https://support.google.com/privacysandbox/answer/15684349?hl=en",
    ])

    footer = doc.sections[0].footer.paragraphs[0]
    footer.alignment = WD_ALIGN_PARAGRAPH.CENTER
    r = footer.add_run("PersonaPulse AI - Comparativo com ferramentas atuais")
    style_run(r, color="777777", size=9)
    doc.save(OUTPUT)


if __name__ == "__main__":
    build()
